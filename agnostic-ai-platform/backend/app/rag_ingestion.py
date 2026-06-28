from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Annotated

from fastapi import APIRouter, File, Form, HTTPException, UploadFile
from pydantic import BaseModel, Field

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".markdown"}
DEFAULT_CHUNK_SIZE = 1000
DEFAULT_CHUNK_OVERLAP = 150

router = APIRouter(prefix="/v1/documents", tags=["RAG Ingestion"])


@dataclass(frozen=True, slots=True)
class LoadedPage:
    text: str
    page: int | None
    source: str
    filename: str


class DocumentChunk(BaseModel):
    text: str
    metadata: dict[str, str | int | None]


class IngestedDocument(BaseModel):
    filename: str
    source: str
    content_type: str | None
    chunks: list[DocumentChunk]
    page_count: int = 0
    chunk_count: int = 0


class IngestionResponse(BaseModel):
    documents: list[IngestedDocument]
    total_documents: int = 0
    total_chunks: int = 0
    chunk_size: int = Field(default=DEFAULT_CHUNK_SIZE, ge=1)
    chunk_overlap: int = Field(default=DEFAULT_CHUNK_OVERLAP, ge=0)


@router.post("/ingest", response_model=IngestionResponse)
async def ingest_documents(
    files: Annotated[list[UploadFile], File(description="Documents to ingest: PDF, DOCX, TXT, MD")],
    chunk_size: Annotated[int, Form(ge=100, le=8000)] = DEFAULT_CHUNK_SIZE,
    chunk_overlap: Annotated[int, Form(ge=0, le=4000)] = DEFAULT_CHUNK_OVERLAP,
) -> IngestionResponse:
    if chunk_overlap >= chunk_size:
        raise HTTPException(status_code=422, detail="chunk_overlap must be smaller than chunk_size")
    if not files:
        raise HTTPException(status_code=422, detail="at least one document is required")

    ingested: list[IngestedDocument] = []
    for upload in files:
        content = await upload.read()
        source = upload.filename or "uploaded-document"
        pages = load_document(content=content, filename=source)
        chunks = chunk_loaded_pages(pages, chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        ingested.append(
            IngestedDocument(
                filename=Path(source).name,
                source=source,
                content_type=upload.content_type,
                chunks=chunks,
                page_count=len({page.page for page in pages if page.page is not None}) or len(pages),
                chunk_count=len(chunks),
            )
        )

    return IngestionResponse(
        documents=ingested,
        total_documents=len(ingested),
        total_chunks=sum(document.chunk_count for document in ingested),
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
    )


def load_document(*, content: bytes, filename: str) -> list[LoadedPage]:
    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise HTTPException(status_code=415, detail=f"unsupported document type {extension!r}; supported: {supported}")

    if extension == ".pdf":
        return _load_pdf(content=content, filename=filename)
    if extension == ".docx":
        return _load_docx(content=content, filename=filename)
    return _load_text(content=content, filename=filename)


def chunk_loaded_pages(
    pages: list[LoadedPage],
    *,
    chunk_size: int = DEFAULT_CHUNK_SIZE,
    chunk_overlap: int = DEFAULT_CHUNK_OVERLAP,
) -> list[DocumentChunk]:
    if chunk_size <= 0:
        raise ValueError("chunk_size must be positive")
    if chunk_overlap < 0 or chunk_overlap >= chunk_size:
        raise ValueError("chunk_overlap must be non-negative and smaller than chunk_size")

    chunks: list[DocumentChunk] = []
    for page in pages:
        text = clean_text(page.text)
        if not text:
            continue
        start = 0
        chunk_index = 0
        while start < len(text):
            end = min(start + chunk_size, len(text))
            chunk_text = text[start:end].strip()
            if chunk_text:
                chunks.append(
                    DocumentChunk(
                        text=chunk_text,
                        metadata={
                            "filename": page.filename,
                            "page": page.page,
                            "source": page.source,
                            "chunk_index": chunk_index,
                            "start_char": start,
                            "end_char": end,
                        },
                    )
                )
                chunk_index += 1
            if end >= len(text):
                break
            start = end - chunk_overlap
    return chunks


def clean_text(text: str) -> str:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    normalized = re.sub(r"[ \t]+", " ", normalized)
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized.strip()


def _load_text(*, content: bytes, filename: str) -> list[LoadedPage]:
    text = content.decode("utf-8", errors="replace")
    return [LoadedPage(text=text, page=1, source=filename, filename=Path(filename).name)]


def _load_pdf(*, content: bytes, filename: str) -> list[LoadedPage]:
    from io import BytesIO

    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise HTTPException(status_code=500, detail="PDF ingestion requires pypdf") from exc

    reader = PdfReader(BytesIO(content))
    pages: list[LoadedPage] = []
    for index, page in enumerate(reader.pages, start=1):
        pages.append(
            LoadedPage(
                text=page.extract_text() or "",
                page=index,
                source=filename,
                filename=Path(filename).name,
            )
        )
    return pages


def _load_docx(*, content: bytes, filename: str) -> list[LoadedPage]:
    from io import BytesIO

    try:
        from docx import Document
    except ImportError as exc:
        raise HTTPException(status_code=500, detail="DOCX ingestion requires python-docx") from exc

    document = Document(BytesIO(content))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            paragraphs.append(" | ".join(cell.text.strip() for cell in row.cells if cell.text.strip()))
    text = "\n".join(paragraphs)
    return [LoadedPage(text=text, page=1, source=filename, filename=Path(filename).name)]

